import xml.etree.ElementTree as tree
import csv
import json
import os
import sys
import spell_fixer

from docx import Document
from itertools import groupby
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Inches
from text_objects import Rectangle, Word, Line, LineType, Column
from text_objects import get_lines, get_columns, get_paragraphs
from text_objects import centre_aligned, merge_words, check_fix_spellings


# print centre aligned text as bold and centred in doc
def print_centre_text(lines, document):
    for line in lines:
        para = document.add_paragraph()
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        para.add_run(line.__str__()).bold = True
        para.add_run("\n")


# print paragraph
def print_paragraph_text(lines, document, page_box, cutoff=0.6, para_margin=20):

    paras = get_paragraphs(lines, para_margin)

    for para_lines in paras:
        first_line = para_lines[0]
        para = document.add_paragraph()
        para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # change font for each paragraph
        style = document.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(10)
        para.style = style

        # make first line bold if it occupies less than width of page
        para.add_run("\n")
        if (first_line.box.x2 - first_line.box.x1)/(page_box.x2 - page_box.x1) < cutoff:
            para.add_run(first_line.__str__()).bold = True
            para.add_run("\n")
            para_lines.pop(0)
        para.add_run(" ".join([line.__str__() for line in para_lines]))
        para.add_run("\n")


# print table
def print_table_text(lines, document, max_col):

    table = document.add_table(rows=len(lines), cols=max_col)
    for line, row in zip(lines, table.rows):
        for word in line:
            row.cells[word.col_num].text = word.value


def filter_and_mark(lines, page_box, page_width, margin=15, cutoff=0.6):

    # get lines with multiple unmerged words
    # these have high chance of forming tables
    for line in lines:
        if len(line.words) > 1 and line.type == -1:
            line.type = LineType.TABLE

    # get centre text
    for line in lines:
        if centre_aligned(line, page_box) and line.type == -1:
            line.type = LineType.CENTRE

    # filter large lines that are likely to be part of paragraphs
    for line in lines:
        if line.type == -1 and (line.box.x2 - line.box.x1)/page_width > cutoff:
            line.type = LineType.PARA

    # get small lines adjacent to table lines even if they have one word
    # these can be empty entries in the table or headings of entries
    # add rest of small lines to paragraph lines
    for i, line in enumerate(lines):

        # first check lower line and assign same type
        if i < len(lines) - 1 and \
            lines[i+1].box.y2 + margin > line.box.y1 and \
                line.type != -1 and lines[i+1].type != -1:
            line.type = lines[i+1].type

        # next check upper line and assign same type
        if i > 0 and lines[i-1].box.y1 - margin < line.box.y2 and \
                line.type != -1 and lines[i-1].type != -1:
            line.type = lines[i-1].type

        # if line is still not assigned a type make it a paragraph line
        if line.type == -1:
            line.type = LineType.PARA


def process_doc(input_path, output_path, parse_mode):
    document = Document()
    xml_file = tree.parse(input_path).getroot()
    for page in xml_file.getchildren():
        page_box = Rectangle(list(map(float, page.attrib["bbox"].split(","))))
        page_mid = (page_box.x1 + page_box.x2)/2
        page_width = page_box.x2 - page_box.x1
        lines = get_lines(page, parse_mode['line_margin'])
        lines = merge_words(lines, parse_mode['merge_margin'])
        lines.sort(reverse=True)

        # filter and mark names with appropriate types
        filter_and_mark(lines, page_box, page_width, parse_mode["adj_margin"], parse_mode["large_cutoff"])

        # correct spellings and discard empty lines
        lines = check_fix_spellings(lines, (LineType.PARA, LineType.TABLE))

        # correct spellings and segmentation for paragraph and table lines
        for line in lines:
            if line.type in [LineType.PARA, LineType.TABLE]:
                for word in line:
                    word.value = spell_fixer.spelling_fixer(word.value)

        # group lines by type and print in docx
        for k, g in groupby(lines, key=lambda x: x.type):
            line_group = list(g)
            if k == LineType.CENTRE:
                print_centre_text(line_group, document)
            elif k == LineType.PARA:
                print_paragraph_text(line_group, document, page_box, parse_mode["large_cutoff"], parse_mode["para_margin"])
            elif k == LineType.TABLE:
                columns = get_columns(line_group)
                columns.sort()
                for i, column in enumerate(columns):
                    for word in column.words:
                        word.col_num = i
                max_col = len(columns)

                # don't make tables for single columns
                if max_col < 2:
                    print_paragraph_text(line_group, document, parse_mode["large_cutoff"], parse_mode["para_margin"])
                else:
                    print_table_text(line_group, document, max_col)

        # create new page break after adding all lines from current page of pdf
        document.add_page_break()

    # change margins
    for section in document.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.2)
        section.right_margin = Inches(0.2)

    document.save(output_path)
