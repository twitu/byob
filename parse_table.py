import xml.etree.ElementTree as tree
from docx import Document
import csv
import json
from enum import Enum
from itertools import groupby
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import spell_fixer
import sys


class LineType(Enum):
    CENTRE = 1
    TABLE = 2
    PARA = 3
    HEADER = 4
    FOOTER = 5


class Rectangle:
    def __init__(self, values):
        self.x1 = values[0]
        self.y1 = values[1]
        self.x2 = values[2]
        self.y2 = values[3]

    # sort in vertically
    def __gt__(self, value):
        return self.y1 > value.y1

    # merge two rectangles to get largest possible rectangles
    def merge(self, value):
        self.x1 = min(self.x1, value.x1)
        self.y1 = min(self.y1, value.y1)
        self.x2 = max(self.x2, value.x2)
        self.y2 = max(self.y2, value.y2)

    def __str__(self):
        return " ".join(map(str, [x1, y1, x2, y2]))


class Word:

    def __init__(self, value, box):
        self.value = value
        self.box = box
        self.row_num = 0
        self.col_num = 0

    def __str__(self):
        return self.value

    # sort from left to right
    def __gt__(self, value):
        return self.box.x1 > self.box.x2

    # concatenate text and merge box
    def merge(self, other):
        self.value += " " + other.value
        self.box.merge(other.box)


class Line:

    def __init__(self, word):
        self.box = word.box
        self.words = [word]
        self.row_num = 0
        self.type = -1

    # add words to list and change bounding box accordingly
    def add_word(self, word):
        self.words.append(word)
        if word.box.y1 < self.box.y1:
            self.box.y1 = word.box.y1
        if self.box.y2 < word.box.y2:
            self.box.y2 = word.box.y2

    # get word from list of words
    def __getitem__(self, item):
        return self.words[item]

    # iterate over list of words
    def __iter__(self):
        return self.words.__iter__()

    # sort lines vertically
    def __gt__(self, value):
        return self.box.y1 > value.box.y1

    # debugging convenience will affect performance
    def __str__(self):
        return " ".join([word.__str__() for word in self.words])


class Column:

    def __init__(self, word):
        self.words = [word]
        self.box = word.box
        self.col_num = 0

    # add word to column and modify bounding box
    def add_word(self, word):
        self.words.append(word)
        if word.box.x1 < self.box.x1:
            self.box.x1 = word.box.x1
        if self.box.x2 < word.box.x2:
            self.box.x2 = word.box.x2

    # get item from list of words
    def __getitem__(self, item):
        return self.words[item]

    # iterate over words in a column
    def __iter__(self):
        return self.words.__iter__()

    # sort columns from left to right
    def __gt__(self, value):
        return self.box.x1 > value.box.x1

    # debugging convenience will affect performance
    def __str__(self):
        return "\n".join([word.__str__() for word in self.words])


class Paragraph:

    def __init__(self, line):
        self.lines = [line]
        self.box = line.box

    # add lines to paragraph and modify bounding box
    def add_line(self, line):
        self.lines.append(line)
        self.box.merge(line.box)

    # get line from list of lines
    def __getitem__(self, item):
        return self.lines[item]

    # iterate over lines in paragraph
    def __iter__(self):
        return self.lines.__iter__()

    # sort vertically
    def __gt__(self, value):
        return self.box.y1 > value.box.x1


# group words that are horizontally in the same line
def get_lines(page_element, margin=5):
    lines = []
    page_box = Rectangle(
        list(map(float, page_element.attrib["bbox"].split(","))))
    for text_box in page_element.getchildren():
        if not text_box.tag == "textbox":
            continue
        for text_line in text_box.getchildren():
            if not text_line.tag == "textline":
                continue
            bbox = Rectangle(
                list(map(float, text_line.attrib["bbox"].split(","))))
            value = ''.join(
                [text_char.text if text_char.text else " " for text_char in text_line.getchildren()])
            value = value.strip()
            if not spell_fixer.spelling_accept(value):
                continue
            value = spell_fixer.spelling_fixer(value)
            word = Word(value, bbox)
            added = False
            for line in lines:
                # check within a margin of error if the text_line can be added in an exisiting line
                if (line.box.y1 - margin < bbox.y1 < line.box.y1 + margin) and (line.box.y2 - margin < bbox.y2 < line.box.y2 + margin):
                    line.add_word(word)
                    added = True

            if not added:
                lines.append(Line(word))

    return lines


# check if box passes centre line and check if it is greater than certain percentage of page width
def filter_centre_word(page_box, word):
    centre = page_box.x2/2
    width = 0.15
    if (word.box.x1 < centre < word.box.x2) and \
            (word.box.x2 - word.box.x1)/(page_box.x2 - page_box.x1) > width:
        return True
    else:
        return False


# take lines of words and create columns from them
def get_columns(lines, margin=20):
    columns = []
    for line in lines:
        for word in line:
            added = False
            for column in columns:
                if column.box.x1 - margin < word.box.x1 < column.box.x1 + margin \
                        or column.box.x2 - margin < word.box.x2 < column.box.x2 + margin \
                        or column.box.x1 + margin < word.box.x1 and word.box.x2 < column.box.x2:
                    column.words.append(word)
                    added = True
            if not added:
                columns.append(Column(word))
    return columns


# merge closely spaced words
def merge_words(lines, margin=15):
    new_lines = []
    for i, line in enumerate(lines):
        line.words.sort()
        merger = line.words[0]
        new_line = Line(merger)
        for i, word in enumerate(line.words[1:]):
            if merger.box.x2 + margin > word.box.x1:
                merger.merge(word)
            else:
                new_line.add_word(merger)
                merger = word
        # avoid adding starting word again
        # happens when line has only one word
        if merger != new_line.words[0]:
            new_line.add_word(merger)
        new_lines.append(new_line)
    return new_lines


def get_paragraphs(lines, margin=5):
    paras = []
    for line in lines:
        added = False
        for para in paras:
            if para.box.y1 - margin < line.box.y2 and line.box.y1 < para.box.y1 or \
                    para.box.y2 + margin > line.box.y1 and line.box.y2 > para.box.y2:
                para.add_line(line)
                added = True

        if not added:
            paras.append(Paragraph(line))

    return paras


def write_to_csv(lines):

    columns = get_columns(lines)
    columns.sort()
    for i, column in enumerate(columns):
        for word in column.words:
            word.col_num = i
    max_col = len(columns)

    with open("table.csv", "w") as csvfile:
        csvwriter = csv.writer(csvfile, delimiter=',',
                               quotechar='"', quoting=csv.QUOTE_ALL)
        for line in lines:
            line.words.sort()
            csv_line = [None]*(max_col)
            for word in line.words:
                csv_line[word.col_num] = word.value
            print(csv_line)
            csvwriter.writerow(csv_line)


def centre_aligned(line, page_box, margin=15, cutoff=0.8):
    mid = (page_box.x1 + page_box.x2)/2
    width = page_box.x2 - page_box.x1

    if line.box.x1 < mid < line.box.x2:
        left = mid - line.box.x1
        right = line.box.x2 - mid
        if (-margin < left - right < margin) and ((line.box.x2 - line.box.x1)/width < cutoff):
            return True

    return False


# function to check for table headings
# table headings takes json file with example headings to compare
def detect_heading(lines, file):
    with open(file) as f:
        data = json.load(f)
    heading_list = data.keys()
    flag = False
    for line in lines:
        for word in line.words:
            for heading in heading_list:
                if word.value.lower() == heading.lower():
                    flag = True
                    break

    return flag


# print centre aligned text as bold and centred in doc
def print_centre_text(lines, document):
    for line in lines:
        para = document.add_paragraph()
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        para.add_run(line.__str__()).bold = True
        para.add_run("\n")


# print paragraph
def print_paragraph_text(lines, document, page_box, cutoff=0.6):
    first_line = lines[0]
    para = document.add_paragraph()
    para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # make first line bold if it occupies less than width of page
    para.add_run("\n")
    if (first_line.box.x2 - first_line.box.x1)/(page_box.x2 - page_box.x1) < cutoff:
        para.add_run(first_line.__str__()).bold = True
        para.add_run("\n")
        lines.pop(0)
    para.add_run(" ".join([line.__str__() for line in lines]))
    para.add_run("\n")


# print table
def print_table_text(lines, document):

    columns = get_columns(lines)
    columns.sort()
    for i, column in enumerate(columns):
        for word in column.words:
            word.col_num = i
    max_col = len(columns)

    table = document.add_table(rows=len(lines), cols=max_col)
    for line, row in zip(lines, table.rows):
        for word in line:
            row.cells[word.col_num].text = word.value


def filter_and_mark(lines, page_box):

    # get lines with multiple unmerged words
    # these have high chance of forming tables
    for line in lines:
        if len(line.words) > 1 and line.type == -1:
            line.type = LineType.TABLE

    # get centre text
    for line in lines:
        if centre_aligned(line, page_box) and line.type == -1:
            line.type = LineType.CENTRE

    # filter large lines that are likely to be part of paragraphs
    for line in lines:
        if (line.box.x2 - line.box.x1)/page_width > 0.6 and line.type == -1:
            line.type = LineType.PARA

    # get small lines adjacent to table lines even if they have one word
    # these can be empty entries in the table or headings of entries
    # add rest of small lines to paragraph lines
    for i, line in enumerate(lines):
        to_add = False
        margin = 15

        # first check lower line and assign same type
        if i < len(lines) - 1 and \
            lines[i+1].box.x2 + margin > line.box.x1 and \
                line.type != -1 and lines[i+1].type != -1:
            line.type = lines[i+1].type

        # next check upper line and assign same type
        if i > 0 and lines[i-1].box.x1 - margin > line.box.x2 and \
                line.type != -1 and lines[i-1].type != -1:
            line.type = lines[i-1].type

        # if line is still not assigned a type make it a paragraph line
        if line.type == -1:
            line.type = LineType.PARA

    # make indiviual table lines as para lines
    for i, line in enumerate(lines):
        if 0 < i < len(lines) - 1 and lines[i].type == LineType.TABLE and \
                lines[i-1].type == lines[i+1].type:
            line.type = LineType.PARA


if __name__ == "__main__":
    document = Document()
    file_name = sys.argv[1]

    xml_file = tree.parse(file_name).getroot()
    for page in xml_file.getchildren():
        page_box = Rectangle(list(map(float, page.attrib["bbox"].split(","))))
        page_mid = (page_box.x1 + page_box.x2)/2
        page_width = page_box.x2 - page_box.x1
        lines = get_lines(page)
        lines = merge_words(lines)
        lines.sort(reverse=True)

        # filter and mark names with appropriate types
        filter_and_mark(lines, page_box)

        # group lines by type and print in docx
        for k, g in groupby(lines, key=lambda x: x.type):
            if k == LineType.CENTRE:
                print_centre_text(list(g), document)
            elif k == LineType.PARA:
                print_paragraph_text(list(g), document, page_box)
            elif k == LineType.TABLE:
                print_table_text(list(g), document)

        # create new page break after adding all lines from current page of pdf
        document.add_page_break()

    document.save(file_name.split(".")[0] + ".docx")
